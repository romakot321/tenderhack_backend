import os
from app.models.dtos import File
import textract
from docx.api import Document
import mimetypes
import requests
import cgi
from urllib.parse import unquote

class FileHandler:

    def __is_task_desciption_document(self, filename):
        KEYWORDS = [
            'тз',
            'техническое задание',
            'техническое_задание',
            'техническое-задание',
        ]
        filename = filename.lower()
        for keyword in KEYWORDS:
            if keyword in filename:
                return True
        return False

    def __is_contract_document(self, filename):
        KEYWORDS = [
            'контракт',
            'договор',
        ]
        filename = filename.lower()
        for keyword in KEYWORDS:
            if keyword in filename:
                return True
        return False

    def __file_docx_to_txt(self, path: str):
        return textract.process(path).decode('utf-8')

    def __file_doc_to_txt(self, path: str):
        return textract.process(path).decode('utf-8')

    def __file_pdf_to_txt(self, path: str):
        return textract.process(path).decode('utf-8')

    def __optimize_text(self, text):
        while ("\n\n" in text):
            text = text.replace("\n\n", "\n")
        return text

    def __extract_tables_docx(self, path):
        document = Document(path)
        if len(document.tables) == 0:
            return None
        else:
            tables = []
            for table in document.tables:
                table_data = []
                keys = None
                for i, row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)
                    if i == 0:
                        keys = tuple(text)
                        continue
                    row_data = dict(zip(keys, text))
                    table_data.append(row_data)
                tables.append(table_data)
            return tables


    def __extract_tables(self, path):
        tables = []
        if path[-5:] == ".docx":
            tables = self.__extract_tables_docx(path)
        return tables

    def __file_to_text(self, path: str):
        text = ""
        if path[-4:] == ".pdf":
            text = self.__file_pdf_to_txt(path)
        elif path[-4:] == ".doc":
            text = self.__file_doc_to_txt(path)
        elif path[-5:] == ".docx":
            text = self.__file_docx_to_txt(path)

        text = self.__optimize_text(text)
        return text

    def __download_file(self, file_id):
        url = "https://zakupki.mos.ru/newapi/api/FileStorage/Download?id=" + file_id
        try:
            query = requests.get(url)
            content_type = query.headers['Content-Type']

            content_disposition = query.headers['Content-Disposition']
            _, content_disposition_params = cgi.parse_header(content_disposition)

            if 'filename*' in content_disposition_params.keys():
                filename = content_disposition_params['filename*']
                filename = unquote(filename)
            else:
                filename = content_disposition_params['filename']
            print(filename)

            is_TZ = self.__is_task_desciption_document(filename)

            file_extension = mimetypes.guess_extension(content_type.split(';')[0])
            output_path = "./assets/session_file"+file_id+file_extension

            response = requests.get(url)
            with open(output_path, 'wb') as file:
                file.write(response.content)
            return output_path, is_TZ
        except:
            print("Error: unable to download file via file_id", file_id)
            return None


    def __delete_file(self, path):
        os.remove(path)

    def __save_txt_file(self, file_id, text, is_TZ):
        path = "./assets/text_file"+file_id+".txt"
        with open(path, 'w') as file:
            file.write(text)
        if is_TZ:
            path_is_TZ = "./assets/bool_tz"+file_id
            with open(path_is_TZ, 'w') as file_is_TZ:
                file_is_TZ.write("")

    def handle_file(self, file_id):
        path_text_file = "./assets/text_file"+file_id+".txt"
        if os.path.exists(path_text_file):
            is_TZ = os.path.exists("./assets/bool_tz"+file_id)
            text = open(path_text_file, 'r').read()
            result = File(
                text=text,
                is_TZ=is_TZ,
            )
            return result
        else:
            path, is_TZ = self.__download_file(file_id)
            if path is None:
                return None

            text = self.__file_to_text(path)

            self.__save_txt_file(file_id, text, is_TZ)

            result = File(
                text=text,
                is_TZ=is_TZ
            )

            self.__delete_file(path)
            return result

handler = FileHandler()
print(handler.handle_file("232508260"))




