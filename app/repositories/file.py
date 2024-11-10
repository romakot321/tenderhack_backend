import os
from app.models.dtos import File
import textract
from docx.api import Document
import mimetypes
import cgi
from urllib.parse import unquote
import asyncio
import aiohttp
import aiofiles

class FileRepository:

    def __is_task_desciption_document(self, filename):
        KEYWORDS = [
            'тз',
            'техническое задание',
            'техническое_задание',
            'техническое-задание',
        ]
        filename = filename.lower()
        for keyword in KEYWORDS:
            if keyword in filename:
                return True
        return False

    def __is_contract_document(self, filename):
        KEYWORDS = [
            'контракт',
            'договор',
        ]
        filename = filename.lower()
        for keyword in KEYWORDS:
            if keyword in filename:
                return True
        return False

    def __file_docx_to_txt(self, path: str):
        return textract.process(path).decode('utf-8')

    def __file_doc_to_txt(self, path: str):
        return textract.process(path).decode('utf-8')

    def __file_pdf_to_txt(self, path: str):
        return textract.process(path).decode('utf-8')

    def __optimize_text(self, text):
        while ("\n\n" in text):
            text = text.replace("\n\n", "\n")
        return text

    def __extract_tables_docx(self, path):
        document = Document(path)
        if len(document.tables) == 0:
            return None
        else:
            tables = []
            for table in document.tables:
                table_data = []
                keys = None
                for i, row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)
                    if i == 0:
                        keys = tuple(text)
                        continue
                    row_data = dict(zip(keys, text))
                    table_data.append(row_data)
                tables.append(table_data)
            return tables


    def __extract_tables(self, path):
        tables = []
        if path[-5:] == ".docx":
            tables = self.__extract_tables_docx(path)
        return tables

    def __file_to_text(self, path: str):
        text = ""
        if path[-4:] == ".pdf":
            text = self.__file_pdf_to_txt(path)
        elif path[-4:] == ".doc":
            text = self.__file_doc_to_txt(path)
        elif path[-5:] == ".docx":
            text = self.__file_docx_to_txt(path)

        text = self.__optimize_text(text)
        return text

    async def __fetch(self, url):
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                data = await response.read()
                headers = response.headers
                return (data, headers)

    async def __write_response_to_file(self, path, data):
        aiofiles_handle = await aiofiles.open(path, 'wb')
        await aiofiles_handle.write(data)

    async def __download_file(self, file_id):
        url = "https://zakupki.mos.ru/newapi/api/FileStorage/Download?id=" + file_id
        data, headers = await self.__fetch(url)

        is_TZ = False
        content_type = headers['Content-Type']
        content_disposition = headers['Content-Disposition']
        _, content_disposition_params = cgi.parse_header(content_disposition)
        if 'filename*' in content_disposition_params.keys():
            filename = content_disposition_params['filename*']
            filename = unquote(filename)
        else:
            filename = content_disposition_params['filename']
        is_TZ = self.__is_task_desciption_document(filename)
        file_extension = mimetypes.guess_extension(content_type.split(';')[0])
        if file_extension is None:
            if filename[-3:] == 'pdf':
                file_extension = '.pdf'
            elif filename[-3:] == 'doc':
                file_extension = '.doc'
            elif filename[-4:] == 'docx':
                file_extension = '.docx'
        output_path = "./assets/session_file" + file_id + file_extension

        await self.__write_response_to_file(output_path, data)

        return output_path, is_TZ



    def __delete_file(self, path):
        os.remove(path)

    async def __save_txt_file(self, file_id, text, is_TZ):
        path = "./assets/text_file" + file_id + ".txt"
        aiofiles_handle_text = await aiofiles.open(path, 'w')
        await aiofiles_handle_text.write(text)
        if is_TZ:
            path_is_TZ = "./assets/bool_tz"+file_id
            aiofiles_handle_bool = await aiofiles.open(path_is_TZ, 'w')
            await aiofiles_handle_bool.write('')

    async def __open_existing_txt_file(self, path):
        async with aiofiles.open(path, mode='r') as file:
            contents = await file.read()
            return contents

    async def handle_file(self, file_id):
        file_id = str(file_id)
        path_text_file = "./assets/text_file"+file_id+".txt"

        if os.path.exists(path_text_file):
            is_TZ = os.path.exists("./assets/bool_tz"+file_id)
            text = await self.__open_existing_txt_file(path_text_file)
            result = File(
                path=path_text_file,
                is_TZ=is_TZ,
            )
            return result
        else:
            download_result = await self.__download_file(file_id)
            if download_result is None:
                return None
            path, is_TZ = download_result

            text = self.__file_to_text(path)

            await self.__save_txt_file(file_id, text, is_TZ)

            result = File(
                path=path_text_file,
                is_TZ=is_TZ
            )

            self.__delete_file(path)
            return result







